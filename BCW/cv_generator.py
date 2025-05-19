from docx import Document

# Data definitions
contact_info = {
    'name': 'Muhammad Faizan Shurjeel',
    'location': 'Lahore, Punjab, Pakistan',
    'phone': '+92-336-7865823',
    'github': 'https://github.com/Faizan-Shurjeel'  # Placeholder; replace with actual GitHub URL
}

objective = 'To secure an engineering position in a reputed company where I can apply my skills in software development, embedded systems, and emerging technologies.'

professional_summary = (
    'A motivated computer engineering graduate with a robust skill set in software development, embedded systems, and emerging technologies. '
    'Experienced in delivering innovative projects using modern programming languages and tools, with a proven ability to adapt and excel in technical challenges. '
    'Passionate about engineering solutions and contributing to team success.'
)

education = {
    'university': 'COMSATS University Islamabad, Lahore Campus',
    'degree': 'Bachelor of Science in Computer Engineering',
    'cgpa': '2.25/4.0',
    'coursework': 'Operating Systems; Databases; Data Structures and Algorithms; Programming Languages; Computer Architecture; Engineering Entrepreneurship; Artificial Intelligence; Computer Networks; Mobile Application Development'
}

skills = {
    'programming_languages': 'C++, C, Java, JavaScript, HTML, CSS, React JS, Dart, Express JS, Flutter, SQL Server, Embedded C/C++, Python',
    'tools_technologies': 'Visual Studio, GitHub, Zrok, Altera Quartus, Proteus, AutoCAD, Jupyter Notebook, Wireshark, Cisco Packet Tracer, MATLAB, GNU/Linux, Android Studio, Firebase, Supabase, PostgreSQL, Vue JS'
}

additional_experiences = [
    'Team Head of Lightroom, Robotics & Automation Society (RAS)',
    'Organized activities for the Association for Computing Machinery (ACM)',
    'Contributed to community development through welfare work for Fruitful Foundation NGO',
    'Participated in mock interviews at Arbisoft',
    'Completed online skill courses on DigiSkills',
    'Active member of the Web3 decentralized tech community',
    'Involved in student club activities at CAC',
    'Learned basics of game development at Mindstorm'
]

chronological_projects = [
    {
        'title': 'Wi-Fi Controlled RC Car',
        'date': '2024',
        'description': [
            'Designed and built a WLAN-controlled car using ESP32.',
            'Competed in IEEE RAS Robo Wars 2024, demonstrating practical engineering skills.'
        ]
    },
    {
        'title': 'Fretch',
        'date': '2023',
        'description': [
            'Developed a Flutter + Rust video downloader application.',
            'Implemented a Rust backend with an HTTP server and integrated yt-dlp for media downloading.',
            'GitHub Repo: [Link]'
        ]
    },
    {
        'title': 'Pak-Blood-Donors',
        'date': '2023',
        'description': [
            'Built a cross-platform mobile app using Flutter to manage blood donors\' records.',
            'Utilized Supabase for backend services and designed a responsive UI with Cupertino Design.',
            'GitHub Repo: [Link]'
        ]
    },
    {
        'title': 'Packet Sniffer GUI',
        'date': '2022',
        'description': [
            'Created a user-friendly GUI for a packet sniffer using thutionlibrary.',
            'Enabled customizable options like interface selection, packet count limits, and protocol filtering.',
            'Displayed sniffed packets in real-time within the GUI.',
            'GitHub Repo: [Link]'
        ]
    },
    {
        'title': 'GUI Calculator in Pure Rust',
        'date': '2022',
        'description': [
            'Developed a GUI calculator application using the Druid framework in Rust.',
            'Provided an interactive interface for basic arithmetic operations.',
            'GitHub Repo: [Link]'
        ]
    },
    {
        'title': 'QuizMasterPro-Web',
        'date': '2021',
        'description': [
            'Collaborated on UI design for a web-based MERN stack quiz management system.',
            'GitHub Repo: [Link]'
        ]
    },
    {
        'title': 'Blockchain Experiment',
        'date': '2021',
        'description': [
            'Built a decentralized application using Golang and Ethereum, exploring blockchain technology.',
            'GitHub Repo: [Link]'
        ]
    }
]

functional_skill_areas = {
    'Software Development': [
        'Built a GUI calculator in Rust using the Druid framework, enabling efficient arithmetic operations.',
        'Developed a packet sniffer GUI with thutionlibrary, offering real-time packet display and customizable filtering options.',
        'Created a cross-platform mobile app, Pak-Blood-Donors, using Flutter and Supabase, streamlining access to donor records.',
        'Designed Fretch, a Flutter + Rust video downloader, integrating yt-dlp and a Rust HTTP server for seamless functionality.'
    ],
    'Embedded Systems': [
        'Engineered a Wi-Fi controlled RC car with ESP32, successfully competing in IEEE RAS Robo Wars 2024.'
    ],
    'Web Development': [
        'Contributed to the UI design of QuizMasterPro-Web, a MERN stack quiz management system, enhancing user interaction.'
    ],
    'Blockchain and Emerging Technologies': [
        'Developed a decentralized application using Golang and Ethereum, demonstrating proficiency in blockchain technology.'
    ]
}

# Functions to add sections to the document
def add_contact_info(doc, contact):
    doc.add_heading(contact['name'], level=1)
    doc.add_paragraph(contact['location'])
    doc.add_paragraph(contact['phone'])
    doc.add_paragraph(f"GitHub: {contact['github']}")

def add_objective(doc, objective):
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(objective)

def add_professional_summary(doc, summary):
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(summary)

def add_education(doc, education):
    doc.add_heading('Education', level=2)
    p = doc.add_paragraph()
    p.add_run(education['university']).bold = True
    doc.add_paragraph(education['degree'])
    doc.add_paragraph(f"CGPA: {education['cgpa']}")
    p = doc.add_paragraph()
    run = p.add_run('Undergraduate Coursework:')
    run.italic = True
    p.add_run(f" {education['coursework']}")

def add_skills(doc, skills):
    doc.add_heading('Skills', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('Programming Languages:')
    run.italic = True
    p.add_run(f" {skills['programming_languages']}")
    p = doc.add_paragraph('', style='List Bullet')
    run = p.add_run('Tools & Technologies:')
    run.italic = True
    p.add_run(f" {skills['tools_technologies']}")

def add_additional_experiences(doc, experiences):
    doc.add_heading('Additional Experiences and Awards', level=2)
    for exp in experiences:
        doc.add_paragraph(exp, style='List Bullet')

def add_technical_experience_chronological(doc, projects):
    doc.add_heading('Technical Experience', level=2)
    for project in projects:
        doc.add_heading(project['title'], level=3)
        p = doc.add_paragraph()
        run = p.add_run(project['date'])
        run.italic = True
        for desc in project['description']:
            doc.add_paragraph(desc, style='List Bullet')

def add_technical_experience_functional(doc, skill_areas):
    doc.add_heading('Key Skills and Accomplishments', level=2)
    for skill, descriptions in skill_areas.items():
        doc.add_heading(skill, level=3)
        for desc in descriptions:
            doc.add_paragraph(desc, style='List Bullet')

# Generate Chronological CV
def create_chronological_cv():
    doc = Document()
    add_contact_info(doc, contact_info)
    add_objective(doc, objective)
    add_education(doc, education)
    add_technical_experience_chronological(doc, chronological_projects)
    add_skills(doc, skills)
    add_additional_experiences(doc, additional_experiences)
    doc.save('chronological_cv.docx')

# Generate Functional CV
def create_functional_cv():
    doc = Document()
    add_contact_info(doc, contact_info)
    add_professional_summary(doc, professional_summary)
    add_technical_experience_functional(doc, functional_skill_areas)
    add_skills(doc, skills)
    add_education(doc, education)
    add_additional_experiences(doc, additional_experiences)
    doc.save('functional_cv.docx')

# Execute the script
if __name__ == '__main__':
    create_chronological_cv()
    create_functional_cv()